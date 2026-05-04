from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt


BASE_DIR = Path(__file__).resolve().parent.parent
SOURCE_MD = BASE_DIR / "docs" / "relatorio_tecnico.md"
OUTPUT_DOCX = BASE_DIR / "docs" / "relatorio_tecnico.docx"
IMAGE_PATTERN = re.compile(r"!\[(?P<alt>.*?)\]\((?P<path>.*?)\)")


def apply_base_style(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1)


def add_code_block(document: Document, lines: list[str]) -> None:
    for line in lines:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(10)
        paragraph.paragraph_format.left_indent = Inches(0.3)


def add_paragraph_with_bold(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    parts = text.split("**")
    for index, part in enumerate(parts):
        run = paragraph.add_run(part)
        if index % 2 == 1:
            run.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY


def add_markdown_image(document: Document, alt_text: str, image_ref: str) -> None:
    image_path = (SOURCE_MD.parent / image_ref).resolve()
    raster_fallback = image_path.with_suffix(".png")

    if image_path.suffix.lower() == ".svg" and raster_fallback.exists():
        image_path = raster_fallback

    if image_path.suffix.lower() == ".svg":
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run(f"[Figura: {alt_text} - arquivo SVG: {image_ref}]")
        run.italic = True
        return

    if image_path.exists():
        document.add_picture(str(image_path), width=Inches(6.5))
        caption = document.add_paragraph()
        caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = caption.add_run(alt_text)
        run.italic = True
        return

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run(f"[Figura não encontrada: {alt_text}]")
    run.italic = True


def markdown_to_docx() -> None:
    document = Document()
    apply_base_style(document)

    lines = SOURCE_MD.read_text(encoding="utf-8").splitlines()
    in_code_block = False
    code_block_language = ""
    code_lines: list[str] = []

    for raw_line in lines:
        line = raw_line.rstrip()

        if line.startswith("```"):
            if in_code_block:
                if code_block_language != "mermaid":
                    add_code_block(document, code_lines)
                code_lines = []
                in_code_block = False
                code_block_language = ""
            else:
                in_code_block = True
                code_block_language = line.removeprefix("```").strip().lower()
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        if not line.strip():
            document.add_paragraph()
            continue

        image_match = IMAGE_PATTERN.fullmatch(line.strip())
        if image_match:
            add_markdown_image(document, image_match.group("alt"), image_match.group("path"))
            continue

        if line.startswith("# "):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run(line[2:].strip())
            run.bold = True
            run.font.size = Pt(16)
            continue

        if line.startswith("## "):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(line[3:].strip())
            run.bold = True
            run.font.size = Pt(14)
            continue

        if line.startswith("### "):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(line[4:].strip())
            run.bold = True
            run.italic = True
            run.font.size = Pt(12)
            continue

        if line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(line[2:].strip())
            continue

        if line.startswith("1. ") or line.startswith("2. ") or line.startswith("3. ") or line.startswith("4. ") or line.startswith("5. ") or line.startswith("6. ") or line.startswith("7. ") or line.startswith("8. ") or line.startswith("9. "):
            paragraph = document.add_paragraph(style="List Number")
            paragraph.add_run(line.split(". ", 1)[1].strip())
            continue

        add_paragraph_with_bold(document, line)

    document.save(OUTPUT_DOCX)
    print(f"Arquivo gerado em {OUTPUT_DOCX}")


if __name__ == "__main__":
    markdown_to_docx()
